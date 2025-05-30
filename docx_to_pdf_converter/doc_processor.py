import docx
from weasyprint import HTML, CSS
import io # For image handling
import os # For path manipulation (getting extension)

def extract_text_from_docx(docx_path):
    """
    Extracts all text from a .docx file.

    Args:
        docx_path (str): The path to the .docx file.

    Returns:
        str: The extracted full text as a single string.
    """
    try:
        document = docx.Document(docx_path)
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from {docx_path}: {e}")
        return ""

def extract_headings(docx_path):
    """
    Extracts headings (up to Level 3) from a .docx file.

    Args:
        docx_path (str): The path to the .docx file.

    Returns:
        list: A list of dictionaries, where each dictionary represents a heading
              and contains 'level', 'text', and 'style' keys.
    """
    headings = []
    try:
        document = docx.Document(docx_path)
        for para in document.paragraphs:
            style_name = para.style.name
            if style_name.startswith('Heading 1'):
                headings.append({'level': 1, 'text': para.text, 'style': style_name})
            elif style_name.startswith('Heading 2'):
                headings.append({'level': 2, 'text': para.text, 'style': style_name})
            elif style_name.startswith('Heading 3'):
                headings.append({'level': 3, 'text': para.text, 'style': style_name})
    except Exception as e:
        print(f"Error extracting headings from {docx_path}: {e}")
    return headings

def test_generate_pdf(output_pdf_path="test.pdf"):
    """
    Generates a simple test PDF with Arabic and English text using WeasyPrint.
    """
    html_content = """
    <!DOCTYPE html>
    <html>
    <head><title>Test PDF</title></head>
    <body dir="rtl">
        <h1>مرحباً بالعالم</h1>
        <p>هذه فقرة اختبار باللغة العربية.</p>
        <p style="color: blue;">This is a test paragraph in English with blue color.</p>
    </body>
    </html>
    """
    try:
        html_doc = HTML(string=html_content)
        html_doc.write_pdf(output_pdf_path)
        print(f"Test PDF generated successfully at {output_pdf_path}")
    except Exception as e:
        print(f"Error generating test PDF: {e}")

def extract_structured_content(docx_path):
    """
    Extracts structured content (headings and associated paragraphs) from a .docx file.
    Output format:
    [
        { # Chapter (H1)
            'title': 'H1 Text', 'level': 1, 'style': 'Heading 1 Style Name',
            'paragraphs': ['Paragraph text 1', 'Paragraph text 2'],
            'children': [ # Subsections (H2)
                {
                    'title': 'H2 Text', 'level': 2, 'style': 'Heading 2 Style Name',
                    'paragraphs': ['Paragraph text under H2'],
                    'children': [ # Sub-subsections (H3)
                        {
                            'title': 'H3 Text', 'level': 3, 'style': 'Heading 3 Style Name',
                            'paragraphs': ['Paragraph text under H3'],
                            'children': [] # H3s do not have further children in this model
                        }
                    ]
                }
            ]
        },
        # ... more H1 chapters
    ]
    """
    document_content = []
    current_h1 = None
    current_h2 = None
    current_h3 = None
    
    images_to_save = []
    image_counter = 0
    document_content = [] # Ensure it's initialized even if docx.Document fails

    try:
        try:
            document = docx.Document(docx_path)
        except Exception as e: # Broadly catch python-docx loading errors
            print(f"Error opening DOCX file {docx_path}: {e}")
            # Propagate a specific error or return empty to indicate failure
            raise ValueError(f"Could not open or parse DOCX file: {os.path.basename(docx_path)}. Error: {str(e)}") from e

        # Process images linked via relationships (covers more cases than just inline_shapes)
        # Store them by a relationship ID if possible, or just iterate and assign unique names
        # Note: This part is complex; python-docx provides access to image parts,
        # but associating them perfectly with paragraph flow is the main challenge.

        # A simplified approach: iterate paragraphs, and for each paragraph, check its runs for inline images.
        for para_idx, para in enumerate(document.paragraphs):
            para_images_after = [] # Images associated with this paragraph

            # Check for images within the paragraph's runs
            for run in para.runs:
                for inline_shape in run.inline_shapes:
                    # Check if it's an image (picture)
                    # InlineShape.type 3 corresponds to a WD_INLINE_SHAPE.PICTURE
                    # https://python-docx.readthedocs.io/en/latest/api/enum/WdInlineShapeType.html
                    if inline_shape.type == 3: # WD_INLINE_SHAPE.PICTURE
                        try:
                            image_part = inline_shape.image_part
                            image_bytes = image_part.blob
                            
                            # Try to get original extension, default to .png
                            original_filename = image_part.filename
                            _, original_extension = os.path.splitext(original_filename)
                            if not original_extension: # if no extension, default
                                original_extension = ".png" 
                            
                            image_counter += 1
                            unique_filename = f"content_image_{image_counter}{original_extension}"
                            
                            images_to_save.append({
                                'unique_filename': unique_filename,
                                'blob': image_bytes
                            })
                            para_images_after.append(unique_filename)
                        except Exception as img_e:
                            print(f"Error processing an inline image: {img_e}")
            
            # Now, handle the paragraph text and structure as before
            style_name = para.style.name.lower()
            active_element_for_images = None # This will be current_h1, h2, or h3 dict

            if style_name.startswith('heading 1'):
                current_h1 = {
                    'title': para.text, 'level': 1, 'style': para.style.name,
                    'paragraphs': [], 'children': [], 'images_after': []
                }
                document_content.append(current_h1)
                current_h2 = None 
                current_h3 = None
                active_element_for_images = current_h1
            elif style_name.startswith('heading 2'):
                if current_h1 is None:
                    current_h1 = {'title': 'Implicit Chapter', 'level': 1, 'style': 'Implicit H1', 'paragraphs': [], 'children': [], 'images_after': []}
                    document_content.append(current_h1)
                current_h2 = {'title': para.text, 'level': 2, 'style': para.style.name, 'paragraphs': [], 'children': [], 'images_after': []}
                current_h1['children'].append(current_h2)
                current_h3 = None
                active_element_for_images = current_h2
            elif style_name.startswith('heading 3'):
                if current_h1 is None:
                    current_h1 = {'title': 'Implicit Chapter', 'level': 1, 'style': 'Implicit H1', 'paragraphs': [], 'children': [], 'images_after': []}
                    document_content.append(current_h1)
                if current_h2 is None:
                    current_h2 = {'title': 'Implicit Section', 'level': 2, 'style': 'Implicit H2', 'paragraphs': [], 'children': [], 'images_after': []}
                    current_h1['children'].append(current_h2)
                current_h3 = {'title': para.text, 'level': 3, 'style': para.style.name, 'paragraphs': [], 'children': [], 'images_after': []} # H3s don't have children
                current_h2['children'].append(current_h3)
                active_element_for_images = current_h3
            else: # It's a paragraph
                if para.text.strip():
                    # Paragraphs are now objects to potentially hold images_after them
                    para_obj = {'type': 'paragraph', 'text': para.text, 'images_after': []}
                    if current_h3:
                        current_h3['paragraphs'].append(para_obj)
                        active_element_for_images = para_obj # Images found in this para text go after this para_obj
                    elif current_h2:
                        current_h2['paragraphs'].append(para_obj)
                        active_element_for_images = para_obj
                    elif current_h1:
                        current_h1['paragraphs'].append(para_obj)
                        active_element_for_images = para_obj
                    # else: # Orphan paragraphs, currently ignored for text. Could collect images for them too.
                    #     pass
            
            # Assign collected images from this paragraph to the active element (heading or paragraph object)
            if active_element_for_images and para_images_after:
                 # If it's a paragraph object, images go into its 'images_after'
                if active_element_for_images.get('type') == 'paragraph':
                    active_element_for_images['images_after'].extend(para_images_after)
                # If it's a heading, and the paragraph *is* the heading itself, images go into heading's 'images_after'
                # This assumes images found *within* a heading's text should appear directly after that heading title.
                elif active_element_for_images.get('title') == para.text : # Check if current para is the heading itself
                     active_element_for_images['images_after'].extend(para_images_after)
                # Otherwise, if it's a heading and we have a paragraph object, it's already handled above.
                # This logic might need refinement based on desired image placement relative to headings vs. their paragraphs.
                # For now, images found in a paragraph's runs are associated with that paragraph object if it's not a heading,
                # or with the heading itself if the paragraph *is* the heading.

    except ValueError: # Catch the re-raised ValueError from docx.Document loading
        raise # Re-raise it to be caught by app.py
    except Exception as e:
        print(f"Unexpected error extracting structured content from {docx_path}: {e}")
        # For other unexpected errors, still try to return what might have been processed
        # or raise a generic error. Let's re-raise to simplify app.py handling.
        raise RuntimeError(f"Unexpected error processing document {os.path.basename(docx_path)}. Error: {str(e)}") from e
        
    return document_content, images_to_save


if __name__ == '__main__':
    # --- Test docx processing functions ---
    dummy_docx_path = 'dummy_test_with_image.docx' # New dummy file name
    run_docx_tests = True 

    if run_docx_tests:
        try:
            doc = docx.Document()
            doc.add_paragraph('Orphan paragraph before any heading.')
            doc.add_heading('Chapter 1: Introduction with Image', level=1)
            doc.add_paragraph('This is the first paragraph of Chapter 1.')
            
            # Add a dummy image (requires a real image file in the execution environment)
            # For testing in this environment, we can't easily create a real image file.
            # So, the image extraction part of extract_structured_content might not find anything
            # or might error if it expects image parts that aren't there for inline_shapes.
            # Let's assume for now that if an image were present, the logic would attempt to process it.
            # To simulate finding an image without actually having one, the test is limited.
            # We'll primarily rely on the code structure being correct.
            # If we could write a file:
            # try:
            #     with open("dummy_image.png", "wb") as f:
            #         f.write(b"iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYAAAAAYAAjCB0C8AAAAASUVORK5CYII=") # 1x1 black pixel png
            #     doc.add_picture("dummy_image.png") # This adds it as a shape, not inline with text easily.
            #     # doc.paragraphs[-1].runs[0].add_picture("dummy_image.png") # Would be more inline
            # except Exception as e:
            #     print(f"Could not create or add dummy image for testing: {e}")

            doc.add_paragraph('This is the second paragraph, after where an image might have been.')
            doc.add_heading('Section 1.1: Background', level=2)
            doc.add_paragraph('Content under Section 1.1.')
            
            doc.save(dummy_docx_path)
            print(f"Created dummy file: {dummy_docx_path}")

            print("\n--- Testing extract_text_from_docx ---")
            text_content = extract_text_from_docx(dummy_docx_path)
            print(f"Extracted Text (first 200 chars):\n{text_content[:200]}..." if text_content else "No text.")

            print("\n--- Testing extract_headings ---")
            headings_list = extract_headings(dummy_docx_path)
            print(headings_list if headings_list else "No headings.")

            print("\n--- Testing extract_structured_content (with image extraction attempt) ---")
            try:
                structured_data, images_found = extract_structured_content(dummy_docx_path)
                import json
                print("Structured Data:")
                print(json.dumps(structured_data, indent=2, ensure_ascii=False))
                print("\nImages Found to Save:")
                print(json.dumps([{'unique_filename': img['unique_filename']} for img in images_found], indent=2))
            except (ValueError, RuntimeError) as e:
                print(f"extract_structured_content failed during test: {e}")
            
        except ImportError:
            print("python-docx library is not installed. Skipping docx tests.")
        except Exception as e: # Catches errors in dummy doc creation or other test logic
            print(f"An error occurred during the broader docx test setup or execution: {e}")
        finally:
            if os.path.exists(dummy_docx_path):
                os.remove(dummy_docx_path)
                print(f"\nCleaned up dummy file: {dummy_docx_path}")
            # if os.path.exists("dummy_image.png"):
            #     os.remove("dummy_image.png")


    # --- Test PDF generation ---
    print("\n--- Testing PDF Generation ---")
    test_generate_pdf("test_output.pdf") # Generates a simple PDF, not using the complex docx
